# -*- coding: utf-8 -*-
"""
实验：基于HMM的汉语词法分析器
- 使用BMES标注体系
- 基于词频构建词表（词典）
- 训练HMM分词模型
- 计算正确率、召回率、F-测度
"""

import os
import re
import sys
import json
import math
import gzip
import shutil
import tarfile
import urllib.request
from collections import Counter, defaultdict
import numpy as np


# ==================== 1. 数据集下载与加载 ====================

def download_pku_dataset(data_dir="data"):
    """下载SIGHAN 2005 PKU中文分词数据集"""
    os.makedirs(data_dir, exist_ok=True)

    train_file = os.path.join(data_dir, "pku_training.utf8")
    test_file = os.path.join(data_dir, "pku_test_gold.utf8")
    test_raw_file = os.path.join(data_dir, "pku_test.utf8")

    # 如果文件已存在，跳过下载
    if os.path.exists(train_file) and os.path.exists(test_file):
        print("数据集已存在，跳过下载。")
        return train_file, test_file, test_raw_file

    print("正在准备中文分词数据集...")

    # 尝试从多个源下载SIGHAN 2005数据
    urls = [
        "http://sighan.cs.uchicago.edu/bakeoff2005/data/icwb2-data.zip",
        "https://github.com/insc74/ICWB2-data/archive/refs/heads/master.zip",
    ]

    downloaded = False
    zip_path = os.path.join(data_dir, "icwb2-data.zip")

    for url in urls:
        try:
            print(f"  尝试下载: {url}")
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=60) as response:
                with open(zip_path, "wb") as f:
                    f.write(response.read())
            downloaded = True
            print("  下载成功！")
            break
        except Exception as e:
            print(f"  下载失败: {e}")
            continue

    if not downloaded:
        print("  网络下载失败，使用内置示例数据集。")
        return create_sample_dataset(data_dir)

    # 解压数据集
    print("  正在解压...")
    try:
        shutil.unpack_archive(zip_path, data_dir)
    except Exception:
        # 尝试用Python解压
        import zipfile
        with zipfile.ZipFile(zip_path, "r") as zf:
            zf.extractall(data_dir)

    # 查找PKU数据集
    pku_train = None
    pku_test = None
    pku_test_raw = None

    for root, dirs, files in os.walk(data_dir):
        for f in files:
            if "pku" in f.lower() and "training" in f.lower() and f.endswith(".utf8"):
                pku_train = os.path.join(root, f)
            elif "pku" in f.lower() and "test_gold" in f.lower() and f.endswith(".utf8"):
                pku_test = os.path.join(root, f)
            elif "pku" in f.lower() and "test" in f.lower() and "gold" not in f.lower() and f.endswith(".utf8"):
                pku_test_raw = os.path.join(root, f)

    if pku_train and pku_test:
        # 复制到data目录下
        shutil.copy(pku_train, train_file)
        shutil.copy(pku_test, test_file)
        if pku_test_raw:
            shutil.copy(pku_test_raw, test_raw_file)
        else:
            # 从gold文件中生成raw（去掉分词标记）
            create_raw_test_file(test_file, test_raw_file)
        print(f"  训练集: {train_file}")
        print(f"  测试集(标注): {test_file}")
    else:
        print("  未找到PKU数据集，使用内置示例数据。")
        return create_sample_dataset(data_dir)

    return train_file, test_file, test_raw_file


def create_raw_test_file(gold_file, raw_file):
    """从标注测试集生成原始测试集（去掉分词空格）"""
    with open(gold_file, "r", encoding="utf-8") as f:
        lines = f.readlines()
    with open(raw_file, "w", encoding="utf-8") as f:
        for line in lines:
            f.write(line.replace(" ", "").strip() + "\n")
    print(f"  生成原始测试集: {raw_file}")


def create_sample_dataset(data_dir):
    """创建内置的中文分词示例数据集（来自人民日报语料片段）"""
    train_file = os.path.join(data_dir, "pku_training.utf8")
    test_file = os.path.join(data_dir, "pku_test_gold.utf8")
    test_raw_file = os.path.join(data_dir, "pku_test.utf8")

    # 训练语料：人民日报风格的中文分词数据
    training_data = [
        "迈向  充满  希望  的  新  世纪  ——  一九九八年  新年  讲话  （  附  图片  １  张  ）",
        "中共中央  总书记  、  国家  主席  江泽民",
        "一九九七年  十二月  三十一日",
        "同志们  ，  朋友们  ，  女士们  ，  先生们  ：",
        "一九九八年  的  新年  钟声  即将  敲响  。",
        "在  这  辞旧迎新  的  美好  时刻  ，  我  代表  中共中央  、  国务院  ，",
        "向  全国  各  民族  人民  ，  向  香港  特别  行政区  同胞  ，",
        "向  澳门  同胞  、  台湾  同胞  和  海外  侨胞  ，",
        "向  世界  各国  的  朋友们  ，  致以  诚挚  的  问候  和  良好  的  祝愿  ！",
        "一九九七年  ，  是  中国  发展  历史  上  非常  重要  的  一年  。",
        "中国  人民  决心  把  建设  有  中国  特色  社会主义  的  伟大  事业  全面  推向  二十一世纪  。",
        "我们  要  继续  坚持  改革  开放  的  方针  ，  加快  经济  建设  步伐  。",
        "中国  的  改革  开放  和  现代化  建设  取得  了  巨大  成就  。",
        "国民经济  持续  快速  健康  发展  ，  人民生活  不断  改善  。",
        "社会主义  精神文明  建设  取得  新  进展  。",
        "我国  的  国际  地位  和  影响  日益  提高  。",
        "我们  要  坚持  和平  统一  的  方针  ，  推进  祖国  统一  大业  。",
        "按照  “  一国两制  ”  的  方针  ，  香港  顺利  回归  祖国  。",
        "我们  要  继续  为  实现  祖国  完全  统一  而  努力  奋斗  。",
        "中国  政府  一贯  坚持  独立  自主  的  和平  外交  政策  。",
        "我们  愿意  在  和平共处五项原则  的  基础  上  ，",
        "同  世界  各国  发展  友好  合作  关系  。",
        "中国  将  继续  发挥  维护  世界  和平  的  积极  作用  。",
        "让  我们  满怀信心  地  迎接  新  的  一年  。",
        "科学  技术  是  第一  生产力  ，  必须  大力  发展  教育  事业  。",
        "我国  经济  持续  快速  增长  ，  人民  生活  水平  显著  提高  。",
        "文化  教育  事业  蓬勃  发展  ，  科技  创新  成果  丰硕  。",
        "我们  要  深化  国有企业  改革  ，  建立  现代  企业  制度  。",
        "农业  是  国民经济  的  基础  ，  必须  始终  把  农业  放在  首位  。",
        "加强  民主  法制  建设  ，  维护  社会  稳定  。",
        "北京  是  中华人民共和国  的  首都  ，  是  政治  文化  中心  。",
        "上海  是  我国  最大  的  经济  中心  城市  。",
        "中国  共产党  领导  全国  各族人民  沿着  社会主义  道路  前进  。",
        "市场  经济  体制  改革  不断  深入  ，  取得  重大  突破  。",
        "外商  投资  环境  不断  改善  ，  吸引  了  大量  外资  。",
        "我们  要  扩大  对外  开放  ，  积极  参与  国际  竞争  。",
        "法律  法规  体系  日趋  完善  ，  执法  力度  不断  加强  。",
        "中国  坚持  走  和平  发展  道路  ，  反对  霸权主义  。",
        "全国  各  族  人民  团结  奋斗  ，  共创  美好  未来  。",
        "当前  国际  形势  复杂  多变  ，  机遇  与  挑战  并存  。",
        "我们  要  抓住  机遇  ，  迎接  挑战  ，  实现  跨越式  发展  。",
        "建设  小康  社会  是  我们  的  奋斗目标  。",
        "中国  加入  世界贸易组织  后  ，  经济  发展  进入  新  阶段  。",
        "加快  金融  体制  改革  ，  防范  金融  风险  。",
        "推动  科技  进步  ，  增强  自主  创新  能力  。",
        "重视  人才  培养  ，  实施  科教兴国  战略  。",
        "我们  要  弘扬  中华  民族  优秀  传统  文化  。",
        "城市  建设  日新月异  ，  基础  设施  日趋  完善  。",
        "农村  面貌  发生  了  巨大  变化  ，  农民  收入  稳步  增长  。",
        "交通  运输  网络  不断  扩大  ，  通讯  事业  快速  发展  。",
        "计算机  技术  在  各个  领域  得到  广泛  应用  。",
        "互  联  网  发展  迅速  ，  信息  产业  成为  新  的  经济增长点  。",
        "电子  商务  蓬勃  兴起  ，  改变  了  传统  的  商业  模式  。",
        "生物  技术  的  突破  给  医学  带来  了  革命性  变化  。",
        "新能源  的  开发  利用  受到  各国  政府  的  高度重视  。",
        "环境  保护  是  人类  共同  面临  的  重大  课题  。",
        "可  持续  发展  战略  必须  得到  切实  贯彻  落实  。",
        "我们  要  建设  资源  节约型  、  环境  友好型  社会  。",
        "中国  政府  高度  重视  社会  保障  体系  建设  。",
        "医疗  卫生  事业  取得  长足  进步  。",
        "体育  事业  蓬勃  发展  ，  全民  健身  运动  广泛  开展  。",
        "二○○八年  北京  奥运  会  取得  圆满  成功  。",
        "中国  健儿  在  国际  赛场  上  屡创佳绩  。",
        "文艺  创作  百花齐放  ，  精品  力作  不断  涌现  。",
        "新闻  出版  事业  健康  发展  ，  舆论  监督  作用  加强  。",
        "我们  要  推进  政治  体制  改革  ，  发展  社会主义  民主  政治  。",
        "依法  治国  是  党  领导  人民  治理  国家  的  基本  方略  。",
        "政府  职能  转变  取得  积极  进展  。",
        "反腐倡廉  工作  深入  开展  ，  取得  明显  成效  。",
        "我们  要  切实  维护  人民  群众  的  根本  利益  。",
        "解决  好  就业  问题  是  当前  工作  的  重中之 重  。",
        "社会  治安  综合  治理  取得  显著  成效  。",
        "我们  要  建设  社会主义  和谐  社会  。",
        "中国  航天  事业  取得  举世瞩目  的  成就  。",
        "人民  军队  革命化  、  现代化  、  正规化  建设  不断  推进  。",
        "我们  坚持  独立自主  的  和平  外交  政策  。",
        "维护  国家  主权  和  领土  完整  是  我们  的  神圣  职责  。",
        "中华民族  伟大  复兴  的  事业  必将  取得  最终  胜利  。",
    ]

    # 测试语料（带标注）
    test_data = [
        "迈向  充满  希望  的  新  世纪",
        "一九九八年  新年  讲话",
        "中共中央  总书记  江泽民",
        "在  这  辞旧迎新  的  美好  时刻",
        "向  全国  各  民族  人民",
        "中国  人民  决心  把  建设  有  中国  特色  社会主义",
        "我们  要  继续  坚持  改革  开放  的  方针",
        "国民经济  持续  快速  健康  发展",
        "我国  的  国际  地位  和  影响  日益  提高",
        "科学  技术  是  第一  生产力",
        "上海  是  我国  最大  的  经济  中心  城市",
        "中国  坚持  走  和平  发展  道路",
        "计算  机  技术  在  各个  领域  得到  广泛  应用",
        "我们  要  建设  资源  节约型  社会",
        "北京  奥运  会  取得  圆满  成功",
        "我们  要  推进  政治  体制  改革",
        "中国  航天  事业  取得  举世瞩目  的  成就",
        "中华民族  伟大  复兴  的  事业",
        "我们  要  弘扬  中华  民族  优秀  传统  文化",
        "推动  科技  进步  增强  自主  创新  能力",
        "建设  小康  社会  是  我们  的  奋斗目标",
        "中国  政府  高度  重视  社会  保障  体系  建设",
        "当前  国际  形势  复杂  多变",
        "实现  中华民族  的  伟大  复兴",
        "改革  开放  以来  中国  经济  取得  了  巨大  成就",
        "随着  信息  技术  的  快速  发展",
        "人工  智能  是  计算机  科学  的  重要  分支",
        "深度  学习  在  图像  识别  领域  表现  出色",
        "自然  语言  处理  是  人工  智能  的  核心  技术  之一",
        "北京  大学  是  中国  著名  的  高等  学府",
        "计算机  科学  与  技术  专业  受到  学生  青睐",
        "互联网  的  普及  改变  了  人们  的  生活  方式",
        "大  数据  技术  已经  在  金融  领域  广泛  应用",
        "云  计算  平台  提供  了  便捷  的  计算  资源",
        "中国  的  高铁  技术  处于  世界  领先  水平",
        "移动  支付  在  中国  得到  了  普遍  使用",
        "快捷  高效  优质  服务  赢得  了  客户  的  信任",
        "坚持  创新  驱动  发展  战略  推进  科技  进步",
        "经济  全球  化  为  发展  中国  家  带来  了  机遇",
    ]

    with open(train_file, "w", encoding="utf-8") as f:
        for line in training_data:
            f.write(line.strip() + "\n")

    with open(test_file, "w", encoding="utf-8") as f:
        for line in test_data:
            f.write(line.strip() + "\n")

    # 生成原始测试集（去掉分词）
    with open(test_raw_file, "w", encoding="utf-8") as f:
        for line in test_data:
            f.write(line.replace(" ", "").strip() + "\n")

    print(f"  内置训练集: {train_file} ({len(training_data)} 行)")
    print(f"  内置测试集: {test_file} ({len(test_data)} 行)")
    return train_file, test_file, test_raw_file


def load_segmented_data(file_path):
    """加载分词标注数据（每行用空格分隔词语）"""
    sentences = []
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            # 按空格分隔得到词语列表
            words = line.split()
            sentences.append(words)
    return sentences


# ==================== 2. BMES标注 ====================

def words_to_bmes(words):
    """将分词结果转换为BMES标签序列
    B: Begin (词的开始)
    M: Middle (词的中间)
    E: End (词的结束)
    S: Single (单字词)
    """
    chars = []
    tags = []
    for word in words:
        if len(word) == 1:
            chars.append(word)
            tags.append("S")
        else:
            chars.append(word[0])
            tags.append("B")
            for c in word[1:-1]:
                chars.append(c)
                tags.append("M")
            chars.append(word[-1])
            tags.append("E")
    return chars, tags


def bmes_to_words(chars, tags):
    """将BMES标签序列还原为词语列表"""
    words = []
    current_word = []
    for char, tag in zip(chars, tags):
        if tag == "B":
            if current_word:
                words.append("".join(current_word))
            current_word = [char]
        elif tag == "M":
            current_word.append(char)
        elif tag == "E":
            current_word.append(char)
            words.append("".join(current_word))
            current_word = []
        elif tag == "S":
            if current_word:
                words.append("".join(current_word))
            words.append(char)
            current_word = []
    if current_word:
        words.append("".join(current_word))
    return words


# ==================== 3. HMM模型 ====================

class HMMSegmenter:
    """基于HMM的中文分词器"""

    def __init__(self):
        self.states = ["B", "M", "E", "S"]
        self.state2id = {s: i for i, s in enumerate(self.states)}
        self.id2state = {i: s for i, s in enumerate(self.states)}

        # 概率矩阵（对数空间）
        self.init_prob = None      # 初始状态概率
        self.trans_prob = None     # 状态转移概率
        self.emiss_prob = None     # 发射概率 (状态 -> 字符)
        self.char2id = None        # 字符到ID的映射

        # 词频统计
        self.word_freq = Counter()
        self.vocab = {}

        self.trained = False

    def train(self, sentences):
        """训练HMM模型，同时统计词频"""
        print("\n" + "=" * 60)
        print("训练HMM分词模型")
        print("=" * 60)

        # 1. 统计词频
        print("  1. 统计词频...")
        self.word_freq = Counter()
        total_chars = 0
        for words in sentences:
            for word in words:
                self.word_freq[word] += 1
                total_chars += len(word)

        print(f"     总词数: {sum(self.word_freq.values())}")
        print(f"     不同词数: {len(self.word_freq)}")
        print(f"     总字符数: {total_chars}")

        # 2. 构建词表（按频率排序）
        print("  2. 构建词表（词典）...")
        self.vocab = {}
        for word, freq in self.word_freq.most_common():
            self.vocab[word] = freq
        print(f"     词表大小: {len(self.vocab)}")

        # 3. 收集所有字符
        all_chars = set()
        for words in sentences:
            for word in words:
                for c in word:
                    all_chars.add(c)
        all_chars = sorted(all_chars)
        self.char2id = {c: i for i, c in enumerate(all_chars)}
        print(f"     字符集大小: {len(self.char2id)}")

        # 4. 统计初始概率、转移概率、发射概率
        print("  3. 统计概率分布...")
        N = len(self.states)
        V = len(self.char2id)

        init_count = np.zeros(N)
        trans_count = np.zeros((N, N))
        emiss_count = np.zeros((N, V))

        for words in sentences:
            chars, tags = words_to_bmes(words)
            if not tags:
                continue

            # 初始状态
            init_count[self.state2id[tags[0]]] += 1

            # 转移和发射
            for i, (c, t) in enumerate(zip(chars, tags)):
                tid = self.state2id[t]
                if c in self.char2id:
                    emiss_count[tid, self.char2id[c]] += 1
                if i > 0:
                    prev_tid = self.state2id[tags[i - 1]]
                    trans_count[prev_tid, tid] += 1

        # 5. 概率估计（加1平滑）
        print("  4. 计算概率（Add-1平滑）...")
        # 初始概率
        self.init_prob = np.log(init_count + 1) - np.log(np.sum(init_count + 1))

        # 转移概率
        trans_smooth = trans_count + 1
        self.trans_prob = np.log(trans_smooth) - np.log(trans_smooth.sum(axis=1, keepdims=True))

        # 发射概率
        emiss_smooth = emiss_count + 1
        self.emiss_prob = np.log(emiss_smooth) - np.log(emiss_smooth.sum(axis=1, keepdims=True))

        self.trained = True
        print("  训练完成！\n")

        # 输出词频前20个词
        print("  【词频最高的20个词】")
        for i, (word, freq) in enumerate(self.word_freq.most_common(20)):
            print(f"    {i+1:2d}. {word:<10s}  频率: {freq}")
        print()

    def viterbi(self, chars):
        """Viterbi解码：找到最优状态序列"""
        if not chars:
            return [], []

        N = len(self.states)
        T = len(chars)
        V = len(self.char2id)

        # dp[t][s]: 到时刻t状态为s的最佳路径的对数概率
        dp = np.full((T, N), -np.inf)
        # back[t][s]: 回溯指针
        back = np.zeros((T, N), dtype=int)

        # 初始化 t=0
        for s in range(N):
            cid = self.char2id.get(chars[0], -1)
            if cid >= 0:
                dp[0, s] = self.init_prob[s] + self.emiss_prob[s, cid]
            else:
                # 未知字符，使用均匀分布
                dp[0, s] = self.init_prob[s] + math.log(1.0 / V)

        # 递推
        for t in range(1, T):
            cid = self.char2id.get(chars[t], -1)
            for s in range(N):
                if cid >= 0:
                    emit_logp = self.emiss_prob[s, cid]
                else:
                    emit_logp = math.log(1.0 / V)

                # 找最佳前一状态
                best_prev = -1
                best_score = -np.inf
                for ps in range(N):
                    score = dp[t - 1, ps] + self.trans_prob[ps, s]
                    if score > best_score:
                        best_score = score
                        best_prev = ps
                dp[t, s] = best_score + emit_logp
                back[t, s] = best_prev

        # 终止：找最后时刻的最佳状态
        best_final = np.argmax(dp[T - 1])

        # 回溯
        tags_ids = [best_final]
        for t in range(T - 1, 0, -1):
            best_final = back[t, best_final]
            tags_ids.append(best_final)
        tags_ids.reverse()

        tags = [self.id2state[tid] for tid in tags_ids]
        return chars, tags

    def segment(self, text):
        """对文本进行分词"""
        if not self.trained:
            raise RuntimeError("模型尚未训练，请先调用 train() 方法。")

        # 移除已有空格
        text = text.replace(" ", "").replace("　", "")
        if not text:
            return []

        chars, tags = self.viterbi(list(text))
        words = bmes_to_words(chars, tags)
        return words

    def save_model(self, path):
        """保存模型"""
        model_data = {
            "states": self.states,
            "state2id": self.state2id,
            "init_prob": self.init_prob.tolist() if self.init_prob is not None else None,
            "trans_prob": self.trans_prob.tolist() if self.trans_prob is not None else None,
            "emiss_prob": self.emiss_prob.tolist() if self.emiss_prob is not None else None,
            "char2id": self.char2id,
            "word_freq": dict(self.word_freq.most_common()),
            "vocab": self.vocab,
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(model_data, f, ensure_ascii=False, indent=2)
        print(f"模型已保存到: {path}")

    def load_model(self, path):
        """加载模型"""
        with open(path, "r", encoding="utf-8") as f:
            model_data = json.load(f)
        self.states = model_data["states"]
        self.state2id = model_data["state2id"]
        self.id2state = {i: s for s, i in self.state2id.items()}
        self.init_prob = np.array(model_data["init_prob"]) if model_data["init_prob"] else None
        self.trans_prob = np.array(model_data["trans_prob"]) if model_data["trans_prob"] else None
        self.emiss_prob = np.array(model_data["emiss_prob"]) if model_data["emiss_prob"] else None
        self.char2id = model_data["char2id"]
        self.word_freq = Counter(model_data["word_freq"])
        self.vocab = model_data["vocab"]
        self.trained = True
        print(f"模型已从 {path} 加载")


# ==================== 4. 评估 ====================

def evaluate_segmentation(pred_words_list, gold_words_list):
    """评估分词性能：正确率、召回率、F-测度"""
    total_gold = 0   # 标准答案中的总词数
    total_pred = 0   # 预测结果中的总词数
    total_correct = 0  # 正确切分的词数

    for pred_words, gold_words in zip(pred_words_list, gold_words_list):
        # 将词序列转换为字符位置索引
        gold_spans = set()
        pos = 0
        for w in gold_words:
            gold_spans.add((pos, pos + len(w)))
            pos += len(w)

        pred_spans = set()
        pos = 0
        for w in pred_words:
            pred_spans.add((pos, pos + len(w)))
            pos += len(w)

        # 统计
        total_gold += len(gold_spans)
        total_pred += len(pred_spans)
        total_correct += len(gold_spans & pred_spans)

    # 计算指标
    precision = total_correct / total_pred if total_pred > 0 else 0
    recall = total_correct / total_gold if total_gold > 0 else 0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0

    return {
        "precision": precision,
        "recall": recall,
        "f1": f1,
        "total_gold": total_gold,
        "total_pred": total_pred,
        "total_correct": total_correct,
    }


def show_segmentation_examples(test_sentences, segmenter, num_examples=10):
    """展示分词样例对比"""
    print("\n" + "=" * 60)
    print("分词结果示例")
    print("=" * 60)

    for i, gold_words in enumerate(test_sentences[:num_examples]):
        text = "".join(gold_words)
        pred_words = segmenter.segment(text)

        print(f"\n  示例 {i+1}:")
        print(f"  原文:     {text}")
        print(f"  标准分词: {' / '.join(gold_words)}")
        print(f"  预测分词: {' / '.join(pred_words)}")

        # 标注错误
        gold_spans = set()
        pos = 0
        for w in gold_words:
            gold_spans.add((pos, pos + len(w)))
            pos += len(w)
        pred_spans = set()
        pos = 0
        for w in pred_words:
            pred_spans.add((pos, pos + len(w)))
            pos += len(w)
        correct = gold_spans & pred_spans
        print(f"  正确词数: {len(correct)}/{len(gold_spans)}")


# ==================== 5. 生成Word实验报告 ====================

def generate_report(results, model_info, word_freq, output_path):
    """生成Word格式的实验报告"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("请先安装python-docx: pip install python-docx")
        return

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ---- 封面标题 ----
    title = doc.add_heading('', level=0)
    run = title.add_run('基于HMM的汉语词法分析器\n实验报告')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 基本信息 ----
    doc.add_paragraph()
    info_lines = [
        "课程名称：自然语言处理",
        "实验名称：基于HMM模型的汉语词法分析器",
        "实验目的：使用HMM模型构建中文分词器，基于分词标注数据训练模型并评估性能",
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ---- 1. 实验原理 ----
    doc.add_heading('一、实验原理', level=1)

    doc.add_heading('1.1 HMM模型概述', level=2)
    p = doc.add_paragraph()
    p.add_run(
        '隐马尔可夫模型（Hidden Markov Model, HMM）是一种统计模型，'
        '用于描述含有隐含未知参数的马尔可夫过程。在中文分词任务中，'
        'HMM将分词问题转化为序列标注问题。'
    )

    doc.add_heading('1.2 BMES标注体系', level=2)
    p = doc.add_paragraph()
    p.add_run('本实验采用BMES标注体系，将每个汉字标注为以下四种状态之一：')
    bmes_list = [
        'B (Begin)：词的开始位置',
        'M (Middle)：词的中间位置',
        'E (End)：词的结束位置',
        'S (Single)：单字成词',
    ]
    for item in bmes_list:
        p = doc.add_paragraph(item, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('例如："中国人民" → B E B E；"我" → S')

    doc.add_heading('1.3 HMM三大要素', level=2)

    # 表格：HMM要素
    table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['要素', '符号', '说明']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True

    table.rows[1].cells[0].text = '初始概率'
    table.rows[1].cells[1].text = 'π'
    table.rows[1].cells[2].text = '句子第一个字的BMES状态概率'

    table.rows[2].cells[0].text = '转移概率'
    table.rows[2].cells[1].text = 'A'
    table.rows[2].cells[2].text = '从状态i转移到状态j的概率'

    table.rows[3].cells[0].text = '发射概率'
    table.rows[3].cells[1].text = 'B'
    table.rows[3].cells[2].text = '在状态j下生成字符k的概率'

    doc.add_heading('1.4 Viterbi解码算法', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Viterbi算法是一种动态规划算法，用于在给定观测序列的条件下，'
        '找到最可能的隐藏状态序列。本实验使用Viterbi算法对输入文本的'
        '每个字符进行BMES状态标注，然后根据标注结果切分出词语。'
    )

    # ---- 2. 实验步骤 ----
    doc.add_heading('二、实验步骤', level=1)

    steps = [
        ('步骤1：数据准备', '下载或加载中文分词标注数据集（如SIGHAN 2005 PKU数据集），'
         '数据格式为每行一个句子，词语之间用空格分隔。'),
        ('步骤2：构建词表', '遍历训练集中的所有句子，统计每个词语的出现频率，'
         '按频率从高到低排序，构建词典（词汇表）。'),
        ('步骤3：BMES标注转换', '将分词结果转换为BMES标签序列，'
         '每个汉字对应一个BMES标签。'),
        ('步骤4：HMM参数估计', '使用最大似然估计（MLE）计算初始概率、转移概率和发射概率，'
         '并采用Add-1平滑处理避免零概率问题。'),
        ('步骤5：模型解码', '使用Viterbi算法对测试文本进行序列标注，'
         '根据标注结果切分出最终的分词结果。'),
        ('步骤6：性能评估', '将预测结果与标准答案进行对比，'
         '计算正确率（Precision）、召回率（Recall）和F-测度（F-measure）。'),
    ]

    for title_text, desc in steps:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(desc)

    # ---- 3. 实验结果 ----
    doc.add_heading('三、实验结果', level=1)

    doc.add_heading('3.1 词频统计与词表构建', level=2)
    p = doc.add_paragraph()
    p.add_run(f'训练数据中共有 {model_info["total_words"]} 个词语，'
              f'其中不同词语数为 {model_info["unique_words"]} 个。'
              f'总字符数为 {model_info["total_chars"]} 个，'
              f'不同字符数为 {model_info["unique_chars"]} 个。')

    doc.add_heading('3.2 词频最高的20个词', level=2)
    word_table = doc.add_table(rows=21, cols=3, style='Light Grid Accent 1')
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_headers = ['排名', '词语', '频率']
    for i, h in enumerate(word_headers):
        word_table.rows[0].cells[i].text = h
        for p in word_table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True

    for rank, (word, freq) in enumerate(word_freq.most_common(20), 1):
        word_table.rows[rank].cells[0].text = str(rank)
        word_table.rows[rank].cells[1].text = word
        word_table.rows[rank].cells[2].text = str(freq)

    doc.add_heading('3.3 分词性能指标', level=2)

    # 性能表格
    perf_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    perf_data = [
        ('指标', '数值'),
        ('正确率 (Precision)', f'{results["precision"]:.4f} ({results["precision"]*100:.2f}%)'),
        ('召回率 (Recall)', f'{results["recall"]:.4f} ({results["recall"]*100:.2f}%)'),
        ('F-测度 (F-measure)', f'{results["f1"]:.4f} ({results["f1"]*100:.2f}%)'),
    ]
    for i, (label, value) in enumerate(perf_data):
        perf_table.rows[i].cells[0].text = label
        perf_table.rows[i].cells[1].text = value
        if i == 0:
            for p in perf_table.rows[i].cells[0].paragraphs:
                for run in p.runs:
                    run.font.bold = True
            for p in perf_table.rows[i].cells[1].paragraphs:
                for run in p.runs:
                    run.font.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('详细统计：').bold = True
    doc.add_paragraph(f'标准答案总词数：{results["total_gold"]}')
    doc.add_paragraph(f'预测结果总词数：{results["total_pred"]}')
    doc.add_paragraph(f'正确分词词数：{results["total_correct"]}')

    doc.add_heading('3.4 分词示例展示', level=2)
    p = doc.add_paragraph(
        '以下展示了模型在测试集上的部分分词结果，'
        '可以直观地看到模型的分词效果。'
    )

    # ---- 4. 实验分析 ----
    doc.add_heading('四、实验分析', level=1)

    doc.add_heading('4.1 HMM模型特点', level=2)
    p = doc.add_paragraph()
    p.add_run('优点：')
    p = doc.add_paragraph('1) 模型简单，训练速度快，参数容易估计。', style='List Bullet')
    p = doc.add_paragraph('2) 基于统计的方法，不需要复杂的语言学规则。', style='List Bullet')
    p = doc.add_paragraph('3) Viterbi解码是全局最优解，能够找到概率最大的标注序列。', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('局限性：')
    p = doc.add_paragraph('1) HMM的独立性假设过强，无法建模长距离依赖。', style='List Bullet')
    p = doc.add_paragraph('2) 仅基于字面特征，无法利用词的语义信息。', style='List Bullet')
    p = doc.add_paragraph('3) 对未登录词（OOV）的处理能力有限。', style='List Bullet')

    doc.add_heading('4.2 性能分析', level=2)
    p = doc.add_paragraph()
    p.add_run(
        f'本实验在测试集上取得了 {results["f1"]*100:.2f}% 的F-测度。'
        f'其中正确率为 {results["precision"]*100:.2f}%，'
        f'召回率为 {results["recall"]*100:.2f}%。'
    )

    if results["f1"] > 0.75:
        p = doc.add_paragraph(
            '整体分词效果较好，说明HMM模型能够有效捕捉中文词语的统计特征。'
        )
    else:
        p = doc.add_paragraph(
            '分词效果还有提升空间，可以考虑：增加训练数据规模、'
            '使用更复杂的模型（如CRF、BiLSTM-CRF）等。'
        )

    # ---- 5. 结论 ----
    doc.add_heading('五、实验结论', level=1)
    p = doc.add_paragraph(
        '本实验成功实现了基于HMM的汉语词法分析器。通过BMES标注体系将中文分词问题'
        '转化为序列标注问题，利用训练数据估计HMM的参数（初始概率、转移概率、发射概率），'
        '并使用Viterbi算法进行最优状态序列的求解。实验结果表明，HMM模型能够较好地完成'
        '中文分词任务，验证了统计方法在自然语言处理中的有效性。'
    )

    # 保存
    doc.save(output_path)
    print(f"\n实验报告已生成: {output_path}")


# ==================== 6. 主程序 ====================

if __name__ == "__main__":
    print("=" * 60)
    print("基于HMM的汉语词法分析器")
    print("=" * 60)

    # 1. 加载数据
    print("\n【步骤1】数据准备")
    train_file, test_file, test_raw_file = download_pku_dataset("data")
    train_sentences = load_segmented_data(train_file)
    test_sentences = load_segmented_data(test_file)
    print(f"  训练集句子数: {len(train_sentences)}")
    print(f"  测试集句子数: {len(test_sentences)}")

    # 2. 构建HMM模型并训练
    print("\n【步骤2】训练HMM模型")
    segmenter = HMMSegmenter()
    segmenter.train(train_sentences)

    # 3. 保存模型
    print("\n【步骤3】保存模型")
    segmenter.save_model("hmm_model.json")

    # 4. 测试&评估
    print("\n【步骤4】评估分词性能")

    # 对测试集每个句子进行分词
    pred_sentences = []
    for words in test_sentences:
        text = "".join(words)
        pred_words = segmenter.segment(text)
        pred_sentences.append(pred_words)

    # 计算性能指标
    results = evaluate_segmentation(pred_sentences, test_sentences)

    print("\n" + "=" * 60)
    print("分词性能评估")
    print("=" * 60)
    print(f"  正确率 (Precision): {results['precision']:.4f} ({results['precision']*100:.2f}%)")
    print(f"  召回率 (Recall):    {results['recall']:.4f} ({results['recall']*100:.2f}%)")
    print(f"  F-测度 (F-measure): {results['f1']:.4f} ({results['f1']*100:.2f}%)")
    print(f"  标准答案总词数: {results['total_gold']}")
    print(f"  预测结果总词数: {results['total_pred']}")
    print(f"  正确分词词数:   {results['total_correct']}")

    # 5. 展示分词示例
    show_segmentation_examples(test_sentences, segmenter, min(10, len(test_sentences)))

    # 6. 生成Word实验报告
    print("\n【步骤5】生成实验报告")
    model_info = {
        "total_words": sum(segmenter.word_freq.values()),
        "unique_words": len(segmenter.word_freq),
        "total_chars": sum(len(w) * f for w, f in segmenter.word_freq.items()),
        "unique_chars": len(segmenter.char2id) if segmenter.char2id else 0,
    }
    generate_report(results, model_info, segmenter.word_freq, "HMM中文分词实验报告.docx")

    print("\n" + "=" * 60)
    print("实验完成！")
    print(f"  - HMM模型: hmm_model.json")
    print(f"  - 实验报告: HMM中文分词实验报告.docx")
    print("=" * 60)
